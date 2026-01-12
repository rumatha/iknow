from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Cm, Inches
from docx.enum.section import WD_ORIENT

from thematic import Thematic
from complex_theme import ComplexTheme
import complex_theme_private_collection
import temporary_team_private_collection
from worksheet import Worksheet
from worksheet_line import WorksheetLine
import worksheet_line_private_collection as wsl
import outlay_tree
import math
import job_place_private_collection
import utils

#===================================================================================================
# Common functions.
#===================================================================================================

def set_document_landscape_orientation(d):
    """
    Set document landscape orientation.

    Parameters
    ----------
    d : Document
        Document.
    """

    section = d.sections[-1]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height

#---------------------------------------------------------------------------------------------------

def set_table_columns_widths(t, ws):
    """
    Set table columns widths.

    Parameters
    ----------
    t : Table
        Word table.
    ws : [float]
        Columns widths in inches.
    """

    ws = [Inches(w) for w in ws]

    for row in t.rows:
        for i, w in enumerate(ws):
            row.cells[i].width = w

#---------------------------------------------------------------------------------------------------

def set_table_text_size(t, s):
    """
    Set table text size.

    Parameters
    ----------
    t : Table
        Table.
    s : int
        Size.
    """

    for row in t.rows:
        for cell in row.cells:
            for par in cell.paragraphs:
                for run in par.runs:
                    font = run.font
                    font.size = Pt(s)

#---------------------------------------------------------------------------------------------------

def merge_table_cells_in_row(t, ri, cifrom, cito):
    """
    Merge table cells in row.

    Parameters
    ----------
    t : Table
        Table.
    ri : int
        Row index.
    cifrom : int
        Collumn index from.
    cito : int
        Column index to.
    """

    row = t.rows[ri].cells

    # Merge.
    c = row[cifrom]
    i = cifrom + 1
    while i <= cito:
        c = c.merge(row[i])
        i = i + 1

#---------------------------------------------------------------------------------------------------

def merge_table_cells_in_column(t, ci, rifrom, rito):
    """
    Merge table cells in column.

    Parameters
    ----------
    t : Table
        Table.
    ci : int
        Column index.
    rifrom : int
        Row index from.
    rito : int
        Row index to.
    """

    col = t.columns[ci].cells

    # Merge.
    c = col[rifrom]
    i = rifrom + 1
    while i <= rito:
        c = c.merge(col[i])
        i = i + 1

#===================================================================================================

class GeneratorWord:
    """
    Master for word documents generation.
    """

    """
    document = Document()

    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    #document.add_picture('monty-truth.png', width=Inches(1.25))

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()

    document.save('demo.docx')
    """

    #-----------------------------------------------------------------------------------------------
    # Common methods.
    #-----------------------------------------------------------------------------------------------

    def __init__(self):
        """
        Init document.
        """

        self.doc = Document()
        sections = self.doc.sections

        for section in sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.25)
            section.right_margin = Cm(1.25)

    #-----------------------------------------------------------------------------------------------

    def save(self, name):
        """
        Save document.

        Parameters
        ----------
        name : str
            File name.
        """

        self.doc.save(name)

    #-----------------------------------------------------------------------------------------------

    def add_empty_line(self):
        """
        Add empty line.
        """

        self.doc.add_paragraph('')

    #-----------------------------------------------------------------------------------------------

    def add_empty_lines(self, n):
        """
        Add empty lines.

        Parameters
        ----------
        n : int
            Count.
        """

        for _ in range(n):
            self.add_empty_line()

    #-----------------------------------------------------------------------------------------------

    def add_paragraph(self, text,
                      alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY, is_bold=False, sz=-1):
        """
        Add paragraph.

        Parameters
        ----------
        text : str
            Text.
        alignment : any
            Text alignment.
        is_bold : bool
            Bold check.
        sz : int
            Size

        Returns
        -------
        Paragraph
            Paragraph.
        """

        # Get paragraph, run and font.
        p = self.doc.add_paragraph()
        r = p.add_run(text)
        f = p.style.font

        # Set properties.
        p.alignment = alignment
        r.bold = is_bold
        f.name = 'Times New Roman'
        if sz >= 0:
            f.size = Pt(sz)
        p.paragraph_format.space_after = Pt(0)
        return p

    #-----------------------------------------------------------------------------------------------

    def add_corner_inscription_supplement_to_order(self, n):
        """
        Add inscription to corner of document.
        Inscription contains text about supplement of order.

        Parameters
        ----------
        n : int
            Number of supplement.
        """

        self.add_paragraph(f'Приложение № {n}\n'
                           'к приказу НИЦ «Курчатовский институт»\n'
                           'от «___» ____________ ____ г. № ______',
                           WD_PARAGRAPH_ALIGNMENT.RIGHT)

    #-----------------------------------------------------------------------------------------------

    def add_signature(self, w, add_line_for_signature):
        """
        Add signature.

        Parameters
        ----------
        w : WorksheetLine
            Worksheet line.
        add_line_for_signature : bool
            Add line for signature.
        """

        # Table and its style.
        t = self.doc.add_table(rows=1, cols=5)
        h = t.rows[0].cells

        # Take job title with first big letter
        h[0].text = w.full_job_title_with_job_place(True)

        h[1].text = ''
        if add_line_for_signature:
            h[2].text = '_____________________\n'\
                        '      (подпись)'
        h[3].text = ''
        h[4].text = w.employee.personal.surname_n_p()

        # Cells sizes.
        xs = [Inches(3.0), Inches(0.5), Inches(2.0), Inches(0.5), Inches(1.5)]
        for row in t.rows:
            for i, x in enumerate(xs):
                row.cells[i].width = x

        self.add_empty_line()

    #-----------------------------------------------------------------------------------------------

    def add_signatures(self, ws, add_line_for_signature=True):
        """
        Add signatures.

        Parameters
        ----------
        ws : [WorksheetLine]
            List of worksheet lines.
        add_line_for_signature : bool
            Add line for signature.
        """

        for w in ws:
            self.add_signature(w, add_line_for_signature)

    #-----------------------------------------------------------------------------------------------
    # Order 3188, form, supplement 07.
    #             exec, supplement 01.
    # Technical task methods.
    #-----------------------------------------------------------------------------------------------

    def add_thematic_results_table_with_TRL(self, th, y):
        """
        Add thematic results table with TRL.

        Parameters
        ----------
        th : Thematic
            Thematic.
        y : int
            Year.
        """

        # Calculate rows count.
        rows_count = 0
        for r in th.results:
            if (not r.is_rid) and (r.year in range(y, y + 3)):
                rows_count = rows_count + 1
        if rows_count == 0:
            return

        # Table and its style.
        t = self.doc.add_table(rows=rows_count + 1, cols=4)
        t.style = 'Table Grid'

        # Head.
        h = t.rows[0].cells
        h[0].text = '№ п/п'
        h[1].text = 'Планируемый результат'
        h[2].text = 'Планируемый УГТ'
        h[3].text = 'Год'

        # Write all results.
        i = 0
        for r in th.results:
            if (not r.is_rid) and (r.year in range(y, y + 3)):
                h = t.rows[1 + i].cells
                h[0].text = str(1 + i)
                h[1].text = r.title
                h[2].text = '1'
                h[3].text = str(r.year)
                i = i + 1

        set_table_columns_widths(t, [0.5, 5.0, 0.5, 0.5])

        self.add_empty_line()

    #-----------------------------------------------------------------------------------------------

    def add_thematic_rids_table_with_TRL(self, th, y):
        """
        Add thematic rids table with TRL.

        Parameters
        ----------
        th : Thematic
            Thematic.
        y : int
            Year
        """

        # Calculate rows count.
        rows_count = 0
        for r in th.results:
            if r.is_rid and (r.year in range(y, y + 3)):
                rows_count = rows_count + 1
        if rows_count == 0:
            return

        # Table and its style.
        t = self.doc.add_table(rows=rows_count + 1, cols=4)
        t.style = 'Table Grid'

        # Head.
        h = t.rows[0].cells
        h[0].text = '№ п/п'
        h[1].text = 'Планируемый результат'
        h[2].text = 'Планируемый УГТ'
        h[3].text = 'Год'

        # Write all results.
        i = 0
        for r in th.results:
            if r.is_rid and (r.year in range(y, y + 3)):
                h = t.rows[1 + i].cells
                h[0].text = str(1 + i)
                h[1].text = f'Планируемый вид РИД: {r.rid_type}.\n'\
                            f'Планируемое название РИД: «{r.rid_name}».\n'\
                            f'Планируемый результат: {r.description}.\n'\
                            f'Планируемое использование: {r.rid_type} будет применяться в '\
                            f'качестве инструментального средства для проведения научных '\
                            f'исследований по тематике {th.title}.'
                h[2].text = '3'
                h[3].text = str(r.year)
                i = i + 1

        set_table_columns_widths(t, [0.5, 5.0, 0.5, 0.5])

        self.add_empty_line()

    #-----------------------------------------------------------------------------------------------

    def add_thematic_indicators_table(self, th, y):
        """
        Add thematic indicators table.

        Parameters
        ----------
        th : Thematic
            Thematic.
        y : int
            Year.
        """

        # Table and its style.
        t = self.doc.add_table(rows=5, cols=6)
        t.style = 'Table Grid'

        # Head.
        h = t.rows[0].cells
        h[0].text = '№ п/п'
        h[1].text = 'Показатель (индикатор)'
        h[2].text = 'Единица измерения'
        h[3].text = f'Год {y}'
        h[4].text = f'Год {y + 1}'
        h[5].text = f'Год {y + 2}'

        # Doctors.
        h = t.rows[1].cells
        h[0].text = '1'
        h[1].text = 'Количество защищенных диссертаций на соискание ученой степени доктора наук'
        h[2].text = 'шт'
        for i in range(3):
            h[3 + i].text = f'{th.ind_doctors(y + i)}'

        # Candidates.
        h = t.rows[2].cells
        h[0].text = '2'
        h[1].text = 'Количество защищенных диссертаций на соискание ученой степени кандидата наук'
        h[2].text = 'шт'
        for i in range(3):
            h[3 + i].text = f'{th.ind_candidates(y + i)}'

        # Rids.
        h = t.rows[3].cells
        h[0].text = '3'
        h[1].text = 'Количество полученных результатов интеллектуальной деятельности'
        h[2].text = 'шт'
        for i in range(3):
            h[3 + i].text = f'{th.ind_rids (y + i)}'

        # Publications.
        h = t.rows[4].cells
        h[0].text = '4'
        h[1].text = 'Публикации в журналах, индексируемых в российских и международных '\
                    'информационно-аналитических системах научного цитирования (Российский индекс '\
                    'научного цитирования или публикации в научных журналах, '\
                    'входящих в «Белый список»)'
        h[2].text = 'шт'
        for i in range(3):
            h[3 + i].text = f'{th.ind_publications(y + i)}'

        set_table_columns_widths(t, [0.5, 5.0, 0.5, 0.5, 0.5, 0.5])

        self.add_empty_line()

    #-----------------------------------------------------------------------------------------------

    def add_thematic_characteristics(self, th, number, y):
        """
        Add thematic characteristics.

        Parameters
        ----------
        th : Thematic
            Thematic.
        number : int
            Number in enumeration.
        y : int
            Year.
        """

        self.add_paragraph(f'7.{number}. Тематика исследований {th.title}.')

        # p. 1-4.
        self.add_paragraph(f'1) Цель работы: {th.goal}')
        self.add_paragraph(f'2) Актуальность и новизна работы: {th.actuality}')
        self.add_paragraph(f'3) Ресурсная обеспеченность работы: {th.resources}')
        self.add_paragraph(f'4) Имеющийся научно-технический задел по работе: {th.background}')

        # p. 5. content
        cs = []
        for r in th.results:
            if (not r.is_rid) and (r.year in range(y, y + 3)):
                cs.append(r.content)
        txt = ' '.join(cs)
        self.add_paragraph(f'5) Основное содержание работ: {txt}')

        # p. 6. years
        self.add_paragraph(f'6) Срок выполнения работы: {y}-{y + 2} годы.')

        # p. 7. results.
        rs = []
        for r in th.results:
            if (not r.is_rid) and (r.year in range(y, y + 3)):
                rs.append(r.title)
        txt = ' '.join(rs)
        self.add_paragraph(f'7) Планируемые результаты: {txt}')
        self.add_thematic_results_table_with_TRL(th, y)
        self.add_thematic_rids_table_with_TRL(th, y)

        # p. 8. indicators
        self.add_paragraph('8) Значение целевых индикаторов и показателей')
        self.add_thematic_indicators_table(th, y)

    #-----------------------------------------------------------------------------------------------

    def add_complex_theme_indicators_table(self, cx, y):
        """
        Add complex theme indicators table.

        Parameters
        ----------
        cx : ComplexTheme
            Complex theme.
        y : int
            Year.
        """

        # Name.
        self.add_paragraph('6. Значение показателей, характеризующих качество работ:')

        # Table and its style.
        t = self.doc.add_table(rows=5, cols=6)
        t.style = 'Table Grid'

        # Head.
        h = t.rows[0].cells
        h[0].text = '№ п/п'
        h[1].text = 'Показатель (индикатор)'
        h[2].text = 'Единица измерения'
        h[3].text = f'Год {y}'
        h[4].text = f'Год {y + 1}'
        h[5].text = f'Год {y + 2}'

        # Doctors.
        h = t.rows[1].cells
        h[0].text = '1'
        h[1].text = 'Количество защищенных диссертаций на соискание ученой степени доктора наук'
        h[2].text = 'шт'
        for i in range(3):
            h[3 + i].text = f'{cx.ind_doctors(y + i)}'

        # Candidates.
        h = t.rows[2].cells
        h[0].text = '2'
        h[1].text = 'Количество защищенных диссертаций на соискание ученой степени кандидата наук'
        h[2].text = 'шт'
        for i in range(3):
            h[3 + i].text = f'{cx.ind_candidates(y + i)}'

        # Rids.
        h = t.rows[3].cells
        h[0].text = '3'
        h[1].text = 'Количество полученных результатов интеллектуальной деятельности'
        h[2].text = 'шт'
        for i in range(3):
            h[3 + i].text = f'{cx.ind_rids (y + i)}'

        # Publications.
        h = t.rows[4].cells
        h[0].text = '4'
        h[1].text = 'Публикации в журналах, индексируемых в российских и международных '\
                    'информационно-аналитических системах научного цитирования (Российский индекс '\
                    'научного цитирования или публикации в научных журналах, '\
                    'входящих в «Белый список»)'
        h[2].text = 'шт'
        for i in range(3):
            h[3 + i].text = f'{cx.ind_publications(y + i)}'

        set_table_columns_widths(t, [0.5, 5.0, 0.5, 0.5, 0.5, 0.5])

    #-----------------------------------------------------------------------------------------------

    def add_complex_theme_characteristics(self, cx, y):
        """
        Add complex theme characteristics.

        Parameters
        ----------
        cx : ComplexTheme
            Complex theme.
        y : int
            Year.
        """

        # p. 1-4.
        self.add_paragraph('1. Подтемы комплексной темы:')
        self.add_paragraph(f'2. Цель выполнения НИР: {cx.goal}')
        self.add_paragraph(f'3. Срок выполнения НИР: {y}-{y + 2} годы.')
        self.add_paragraph(f'4. Актуальность НИР: {cx.actuality}')

        # p. 5. Get all results (without RIDs).
        rs = []
        for th in cx.thematics:
            for r in th.results:
                if (not r.is_rid) and (r.year in range(y, y + 3)):
                    rs.append(r.title)
        txt = ' '.join(rs)
        self.add_paragraph(f'5. Планируемые результаты НИР: {txt}')

        # p. 6.
        self.add_complex_theme_indicators_table(cx, y)

        # p. 6.
        self.add_empty_line()
        self.add_paragraph('7. Основное содержание работ по тематикам исследований:')

        # Add characteristics for all thematics.
        for i, th in enumerate(cx.thematics):
            self.add_thematic_characteristics(th, i + 1, y)

    #-----------------------------------------------------------------------------------------------

    @staticmethod
    def generate_form_gos_assignment_3188_07_technical_task(theme, y, out):
        """
        Generate form gos assignment.
        Supplement 7 - technical task.

        Parameters
        ----------
        theme : ComplexTheme
            Complex theme.
        y : int
            Year.
        out : str
            Out file name.
        """

        w = GeneratorWord()

        # Title.
        w.add_paragraph('ТЕХНИЧЕСКОЕ ЗАДАНИЕ\nна выполнение научно-исследовательской работы '
                        f'по комплексной теме {theme.title}',
                        WD_PARAGRAPH_ALIGNMENT.CENTER, True)
        w.add_empty_line()

        # Main information.
        w.add_complex_theme_characteristics(theme, y)
        w.add_empty_line()

        # Signatures and save.
        w.add_signatures([theme.manager, wsl.shabanov_bm])
        w.save(out + '.docx')

    #-----------------------------------------------------------------------------------------------

    @staticmethod
    def generate_exec_gos_assignment_3188_01_technical_task(theme, y, out):
        """
        Generate form gos assignment.
        Supplement 7 - technical task.

        Parameters
        ----------
        theme : ComplexTheme
            Complex theme.
        y : int
            Year.
        out : str
            Out file name.
        """

        w = GeneratorWord()

        # Supplement in the corner.
        w.add_corner_inscription_supplement_to_order(1)
        w.add_empty_line()

        # Title.
        w.add_paragraph('ТЕХНИЧЕСКОЕ ЗАДАНИЕ\nна выполнение научно-исследовательской работы '
                        f'по комплексной теме {theme.title}',
                        WD_PARAGRAPH_ALIGNMENT.CENTER, True)
        w.add_empty_line()

        # Main information.
        w.add_complex_theme_characteristics(theme, y)
        w.add_empty_line()

        # Signatures and save.
        w.add_signatures([theme.manager, wsl.shabanov_bm])
        w.save(out + '.docx')

    #-----------------------------------------------------------------------------------------------
    # Order 3188, form, supplement 08.
    #             exec, supplement 02.
    # Calendar plan methods.
    #-----------------------------------------------------------------------------------------------

    def add_calendar_plan_table_for_3188_form_08(self, theme, y):
        """
        Add calendar plan for 8 supplement.

        Parameters
        ----------
        theme : ComplexTheme
            Complex theme.
        y : int
            Year.
        """

        # rows count
        rows_count = 2 + 4 * len(theme.thematics)

        # Table and its style.
        t = self.doc.add_table(rows=rows_count, cols=13)
        t.style = 'Table Grid'

        # Names of fields.
        h = t.rows[0].cells
        h[0].text = '№ п/п'
        h[1].text = 'Содержание выполняемых работ'
        h[2].text = 'Ожидаемый результат выполнения работ '\
                    '/ Период реализации работ в рамках тематики исследований'
        h[3].text = 'Планируемый к созданию результат интеллектуальной деятельности (далее - '\
                    'РИД) и уровень готовности разрабатываемых или разработанных технологий '\
                    '(далее - УГТ)'
        h[8].text = 'Ожидаемый результат по показателям качества работ, '\
                    'установленных государственным заданием'
        h[10].text = 'Состав отчетной документации'
        h[11].text = 'Ответственный руководитель работ по комплексной теме\n\n '\
                     'Ответственное структурное подразделение Центра за выполнение работ по '\
                     'комплексной теме\n\n Ответственный руководитель работ по подтеме '\
                     'комплексной темы с указанием структурного подразделения\n\n Ответственный '\
                     'руководитель работ по тематике исследований подтемы'
        h[12].text = 'Стоимость, рублей'
        h = t.rows[1].cells
        h[3].text = 'Вид планируемого к созданию РИД'
        h[4].text = 'Планируемое наименование планируемого к созданию РИД'
        h[5].text = 'Краткое описание планируемого к созданию РИД'
        h[6].text = 'Планируемый УГТ'
        h[7].text = 'Срок создания РИД (мм.гггг)'
        h[8].text = 'Ожидаемый результат'
        h[9].text = 'Срок исполнения'

        # Shape.
        merge_table_cells_in_column(t, 0, 0, 1)
        merge_table_cells_in_column(t, 1, 0, 1)
        merge_table_cells_in_column(t, 2, 0, 1)
        merge_table_cells_in_row(t, 0, 3, 7)
        merge_table_cells_in_row(t, 0, 8, 9)
        merge_table_cells_in_column(t, 10, 0, 1)
        merge_table_cells_in_column(t, 11, 0, 1)
        merge_table_cells_in_column(t, 12, 0, 1)

        # Start row number.
        hi = 2

        # Walk all thematics.
        for thematici, thematic in enumerate(theme.thematics):

            # Add thematic title.
            h = t.rows[hi].cells
            h[0].text = f'{thematici + 1}.'
            h[1].text = thematic.title
            merge_table_cells_in_row(t, hi, 1, 12)
            hi = hi + 1

            # Walk all years.
            for year in range(y, y + 3):
                h = t.rows[hi].cells

                # Get common results and rids.
                rs_com = [r for r in thematic.results if (not r.is_rid) and (r.year == year)]
                rs_rid = [r for r in thematic.results if r.is_rid and (r.year == year)]

                # Results and content.
                h[1].text = ' '.join([r.content for r in rs_com])
                h[2].text = ' '.join([r.title for r in rs_com])

                # RID information.
                h[3].text = '\n\n'.join([r.rid_type for r in rs_rid])
                h[4].text = '\n\n'.join([r.rid_name for r in rs_rid])
                h[5].text = '\n\n'.join([r.description for r in rs_rid])
                h[6].text = '\n\n'.join(['Третий УГТ' for _ in rs_rid])
                h[7].text = '\n\n'.join([f'12.{year}' for _ in rs_rid])

                # Results and values.
                txt = f'докторские диссертации - {thematic.ind_doctors(year)};\n'\
                      f'кандидатские диссертации- {thematic.ind_candidates(year)};\n'\
                      f'РИД - {len(rs_rid)};\nпубликации - {thematic.ind_publications(year)}'
                h[8].text = txt
                h[9].text = f'{year} год'

                # Doc, responsible.
                h[10].text = 'Аннотационный отчет - ежеквартально;\n\n'\
                             'Итоговый отчет о НИР.'
                h[11].text = f'{theme.manager.full_name_with_full_job_title_in_brackets()}\n\n'\
                             f'{job_place_private_collection.osspv.name}'

                # Money without hoz spents.
                x = thematic.outlay['II'].xmoney_rubles_with_kopecks_str
                h[12].text = f'{x}'

                # Move row counter.
                hi = hi + 1

        set_table_text_size(t, 8)

    #-----------------------------------------------------------------------------------------------

    def add_calendar_plan_table_for_3188_exec_02(self, theme, y):
        """
        Add calendar plan for 8 supplement.

        Parameters
        ----------
        theme : ComplexTheme
            Complex theme.
        y : int
            Year.
        """

        # rows count
        rows_count = 2 + 4 * len(theme.thematics)

        # Table and its style.
        t = self.doc.add_table(rows=rows_count, cols=10)
        t.style = 'Table Grid'

        # Names of fields.
        h = t.rows[0].cells
        h[0].text = '№ п/п'
        h[1].text = 'Содержание выполняемых работ'
        h[2].text = 'Ожидаемый результат'
        h[7].text = 'Состав отчетной документации'
        h[8].text = 'Ответственный руководитель работ по комплексной теме\n\n '\
                    'Ответственное структурное подразделение Центра за выполнение работ по '\
                    'комплексной теме\n\n Ответственный руководитель работ по подтеме '\
                    'комплексной темы с указанием структурного подразделения\n\n Ответственный '\
                    'руководитель работ по тематике исследований подтемы '\
                    'с указанием структурного подразделения'
        h[9].text = 'Стоимость, рублей'
        h = t.rows[1].cells
        h[2].text = 'Вид планируемого к созданию результата интеллектуальной деятельности '\
                    '(далее - РИД)'
        h[3].text = 'Планируемое наименование планируемого к созданию РИД'
        h[4].text = 'Краткое описание планируемого к созданию РИД'
        h[5].text = 'Планируемый уровень готовности разрабатываемых или разработанных технологий '\
                    '(далее - УГТ)'
        h[6].text = 'Срок создания РИД (мм.гггг)'

        # Shape.
        merge_table_cells_in_column(t, 0, 0, 1)
        merge_table_cells_in_column(t, 1, 0, 1)
        merge_table_cells_in_row(t, 0, 2, 6)
        merge_table_cells_in_column(t, 7, 0, 1)
        merge_table_cells_in_column(t, 8, 0, 1)
        merge_table_cells_in_column(t, 9, 0, 1)

        # Start row number.
        hi = 2

        # Walk all thematics.
        for thematici, thematic in enumerate(theme.thematics):

            # Add thematic title.
            h = t.rows[hi].cells
            h[0].text = f'{thematici + 1}.'
            h[1].text = thematic.title
            merge_table_cells_in_row(t, hi, 1, 9)
            hi = hi + 1

            # Walk all years.
            for year in range(y, y + 3):
                h = t.rows[hi].cells

                # Get common results and rids.
                rs_com = [r for r in thematic.results if (not r.is_rid) and (r.year == year)]
                rs_rid = [r for r in thematic.results if r.is_rid and (r.year == year)]

                # Content.
                txt = ' '.join([r.content for r in rs_com])
                h[1].text = f'{year} год\n{txt}'

                # RID information.
                h[2].text = '\n\n'.join([r.rid_type for r in rs_rid])
                h[3].text = '\n\n'.join([r.rid_name for r in rs_rid])
                h[4].text = '\n\n'.join([r.description for r in rs_rid])
                h[5].text = '\n\n'.join(['Третий УГТ' for _ in rs_rid])
                h[6].text = '\n\n'.join([f'12.{year}' for _ in rs_rid])

                # Doc, responsible.
                h[7].text = 'Аннотационный отчет - ежеквартально;\n\n'\
                            'Итоговый отчет о НИР.'
                h[8].text = f'{theme.manager.full_name_with_full_job_title_in_brackets()}\n\n'\
                            f'{job_place_private_collection.osspv.name}'

                # All money.
                x = thematic.outlay.xmoney_rubles_with_kopecks_str
                h[9].text = f'{x}'

                # Move row counter.
                hi = hi + 1

        set_table_text_size(t, 7)

    #-----------------------------------------------------------------------------------------------

    @staticmethod
    def generate_form_gos_assignment_3188_08_calendar_plan(theme, y, out):
        """
        Generate form gos assignment.
        8 supplement - calendar plan.

        Parameters
        ----------
        theme : ComplexTheme
            Complex theme.
        y : int
            Year.
        out : str
            Output file name.
        """

        w = GeneratorWord()
        set_document_landscape_orientation(w.doc)

        # Title.
        w.add_paragraph(f'КАЛЕНДАРНЫЙ ПЛАН\nна {y} год '
                        f'и плановый период {y + 1} и {y + 2} годов '
                        'на выполнение научно-исследовательской работы '
                        f'по комплексной теме {theme.title}',
                        WD_PARAGRAPH_ALIGNMENT.CENTER, True)
        w.add_empty_line()

        # Add tables.
        w.add_calendar_plan_table_for_3188_form_08(theme, y)
        w.add_empty_line()

        # Add signatures and save.
        w.add_signatures([theme.manager, wsl.shabanov_bm])
        w.save(out + '.docx')

    #-----------------------------------------------------------------------------------------------

    @staticmethod
    def generate_exec_gos_assignment_3188_02_calendar_plan(theme, y, out):
        """
        Generate exec gos assignment.
        2 supplement - calendar plan.

        Parameters
        ----------
        theme : ComplexTheme
            Complex theme.
        y : int
            Year.
        out : str
            Output file name.
        """

        w = GeneratorWord()
        set_document_landscape_orientation(w.doc)

        # Supplement.
        w.add_corner_inscription_supplement_to_order(2)
        w.add_empty_line()

        # Title.
        w.add_paragraph(f'КАЛЕНДАРНЫЙ ПЛАН\nна выполнение научно-исследовательской '
                        f'работы по комплексной теме {theme.title}',
                        WD_PARAGRAPH_ALIGNMENT.CENTER, True)
        w.add_empty_line()

        # Add tables.
        w.add_calendar_plan_table_for_3188_exec_02(theme, y)
        w.add_empty_line()

        # Add signatures and save.
        w.add_signatures([theme.manager, wsl.shabanov_bm])
        w.save(out + '.docx')

    #-----------------------------------------------------------------------------------------------
    # Order 3188, form, supplement 09.
    #             exec, supplement 03.
    # Outlay methods.
    #-----------------------------------------------------------------------------------------------

    def add_outlay_table_theme_with_subthemes(self, outlay, y, hoz):
        """
        Add outlay table for one theme - subthemes.

        NB! We have no subthemes.

        Parameters
        ----------
        outlay : outlay_tree.Nod
            Outlay
        y : int
            Year.
        hoz : bool
            Add hoz spents.
        """

        self.add_paragraph('рублей', WD_PARAGRAPH_ALIGNMENT.RIGHT)

        # Flatten outlay.
        outlay_lines = outlay.flatten()
        k = outlay['II'].count()

        # Table and its style.
        t = self.doc.add_table(rows=1+k, cols=4)
        t.style = 'Table Grid'

        # Add row.
        def add_row(i, h0, h1, h2, h3):
            h = t.rows[i].cells
            h[0].text, h[1].text, h[2].text, h[3].text = h0, h1, h2, h3

        # Add head.
        add_row(0,
                '№ п/п',
                'Наименование статей расходов',
                'Всего стоимость, рублей',
                'В том числе по подтемам комплексной темы')

        # Add all lines.

        rowi = 1

        # If it is form without hoz then print first row.
        if not hoz:
            x = outlay['II'].xmoney_rubles_with_kopecks_str
            add_row(rowi,'', 'ВСЕГО ЗАТРАТ, в том числе:', f'{x}', f'{x}')
            rowi = rowi + 1

        for ol in outlay_lines:

            # Do not print
            if (ol.label == 'I.') or (ol.label == 'II.'):
                continue

            # Do not print hoz spents.
            if (not hoz) and (ol.label == 'III.'):
                break

            x = ol.xmoney_rubles_with_kopecks_str
            add_row(rowi, ol.label, ol.name, f'{x}', f'{x}')
            rowi = rowi + 1

        set_table_columns_widths(t, [0.5, 5.0, 1.5, 1.5])

    #-----------------------------------------------------------------------------------------------

    def add_short_outlay_table(self, cx, y):
        """
        Add short outlay table.

        Parameters
        ----------
        cx : ComplexTheme
            Complex theme.
        y : int
            Year.
        """

        self.add_paragraph('рублей', WD_PARAGRAPH_ALIGNMENT.RIGHT)

        k = len(cx.thematics)

        # Table and its style.
        t = self.doc.add_table(rows=2*k + 4, cols=5)
        t.style = 'Table Grid'

        # Head.
        h = t.rows[0].cells
        h[0].text = '№ п/п'
        h[1].text = 'Наименование статей расходов'
        h[2].text = f'{y} год'
        h[3].text = f'{y + 1} год'
        h[4].text = f'{y + 2} год'

        # 1.
        h = t.rows[1].cells
        h[0].text = '1.'
        h[1].text = 'ВСЕГО по комплексной теме, в том числе'

        # 2.
        h = t.rows[2].cells
        h[0].text = '2.'
        h[1].text = 'Прямые и общепроизводственные затраты, из них:'

        # 3.
        h = t.rows[k + 3].cells
        h[0].text = '3.'
        h[1].text = 'Общехозяйственные расходы, из них:'

        # Process all thematics.
        for i, th in enumerate(cx.thematics):
            #
            h = t.rows[3 + i].cells
            h[0].text = f'2.{i + 1}.'
            h[1].text = f'Тематика исследований {th.title}'
            x = th.outlay['II'].xmoney_rubles_with_kopecks_str
            h[2].text = f'{x}'
            h[3].text = f'{x}'
            h[4].text = f'{x}'
            #
            h = t.rows[4 + k + i].cells
            h[0].text = f'3.{i + 1}.'
            h[1].text = f'Тематика исследований {th.title}'
            x = th.outlay['III'].xmoney_rubles_with_kopecks_str
            h[2].text = f'{x}'
            h[3].text = f'{x}'
            h[4].text = f'{x}'

        # Write sums.
        h = t.rows[2].cells
        direct = cx.outlay['II'].xmoney_rubles_with_kopecks_str
        h[2].text = f'{direct}'
        h[3].text = f'{direct}'
        h[4].text = f'{direct}'
        h = t.rows[6].cells
        hoz = cx.outlay['III'].xmoney_rubles_with_kopecks_str
        h[2].text = f'{hoz}'
        h[3].text = f'{hoz}'
        h[4].text = f'{hoz}'
        h = t.rows[1].cells
        s = cx.outlay.xmoney_rubles_with_kopecks_str
        h[2].text = f'{s}'
        h[3].text = f'{s}'
        h[4].text = f'{s}'

        set_table_columns_widths(t, [0.5, 4.0, 1.5, 1.5, 1.5])

    #-----------------------------------------------------------------------------------------------

    def add_outlay_table(self, outlay, y):
        """
        Add outlay table.

        Parameters
        ----------
        outlay : outlay_tree.Nod
            Outlay.
        y : int
            Year.
        """

        self.add_paragraph('рублей', WD_PARAGRAPH_ALIGNMENT.RIGHT)

        outlay_lines = outlay.flatten()
        k = len(outlay_lines)

        # Table and its style.
        t = self.doc.add_table(rows=1+k, cols=5)
        t.style = 'Table Grid'

        # Head.
        h = t.rows[0].cells
        h[0].text = '№ п/п'
        h[1].text = 'Наименование статей расходов'
        h[2].text = f'{y} год'
        h[3].text = f'{y + 1} год'
        h[4].text = f'{y + 2} год'

        # Add all lines.
        for i in range(k):
            ol = outlay_lines[i]
            h = t.rows[1 + i].cells
            h[0].text = ol.label
            h[1].text = ol.name
            x = ol.xmoney_rubles_with_kopecks_str
            h[2].text = f'{x}'
            h[3].text = f'{x}'
            h[4].text = f'{x}'

        set_table_columns_widths(t, [0.5, 4.0, 1.5, 1.5, 1.5])

    #-----------------------------------------------------------------------------------------------

    @staticmethod
    def generate_form_gos_assignment_3188_09_outlay(cx, y, out):
        """
        Generate
        'form gos assignment, supplement 09 - pre outlay'.

        Parameters
        ----------
        cx : ComplexTheme
            Complex theme.
        y : int
            Year.
        out : str
            Out file.
        """

        # Create document.
        w = GeneratorWord()

        # Title.
        w.add_paragraph('ПРЕДВАРИТЕЛЬНАЯ СМЕТА\n'
                        'прямых и общепроизводственных расходов, непосредственно связанных с '
                        'выполнением научно-исследовательской работы '
                        f'по комплексной теме {cx.title} на очередной {y} год '
                        f'и плановый период {y + 1} и {y + 2} годов',
                        WD_PARAGRAPH_ALIGNMENT.CENTER, True)

        # First table.
        w.add_paragraph('1. Предварительная смета прямых и общепроизводственных расходов, '
                        'непосредственно связанных с выполнением научно-исследовательской работы '
                        f'на очередной {y} год',
                        WD_PARAGRAPH_ALIGNMENT.CENTER, True)
        w.add_outlay_table_theme_with_subthemes(cx.outlay, y, False)
        w.add_empty_line()

        # Second table.
        w.add_paragraph('2. Предварительная смета прямых и общепроизводственных расходов, '
                        'непосредственно связанных с выполнением научно-исследовательской работы '
                        f'на первый плановый год - {y + 1} год',
                        WD_PARAGRAPH_ALIGNMENT.CENTER, True)
        w.add_outlay_table_theme_with_subthemes(cx.outlay, y + 1, False)
        w.add_empty_line()

        # Third table.
        w.add_paragraph('3. Предварительная смета прямых и общепроизводственных расходов, '
                        'непосредственно связанных с выполнением научно-исследовательской работы '
                        f'на второй плановый год - {y + 2} год',
                        WD_PARAGRAPH_ALIGNMENT.CENTER, True)
        w.add_outlay_table_theme_with_subthemes(cx.outlay, y + 2, False)
        w.add_empty_line()

        # Add signatures and close file.
        w.add_signatures([cx.manager, wsl.shabanov_bm, wsl.smirnnova_oe, wsl.petrischev_av])
        w.save(out + '.docx')

    #-----------------------------------------------------------------------------------------------

    @staticmethod
    def generate_exec_gos_assignment_3188_03_outlay(cx, y, out):
        """
        Generate outlay document.

        Parameters
        ----------
        cx : ComplexTheme
            Complex theme.
        y : int
            Start year.
        out : str
            Out file name.
        """

        w = GeneratorWord()

        # Incription.
        w.add_corner_inscription_supplement_to_order(3)
        w.add_empty_line()

        # Title.
        w.add_paragraph('ИТОГОВАЯ СМЕТА\nрасходов на выполнение работы '
                        f'по комплексной теме {cx.title} на очередной {y} год '
                        f'и плановый период {y + 1} и {y + 2} годов',
                        WD_PARAGRAPH_ALIGNMENT.CENTER, True)
        w.add_empty_line()

        # Add short outlay for complex theme.
        w.add_paragraph(f'1. ИТОГОВАЯ СВОДНАЯ СМЕТА\nпо комплексной теме {cx.title}',
                        WD_PARAGRAPH_ALIGNMENT.CENTER, True)
        w.add_short_outlay_table(cx, y)
        w.add_empty_line()

        # Add outlay for complex theme.
        w.add_paragraph('2. ИТОГОВАЯ СМЕТА\n'
                        f'на очередной {y} год и плановый период {y + 1} и {y + 2} годов\n'
                        f'по комплексной теме {cx.title}',
                        WD_PARAGRAPH_ALIGNMENT.CENTER, True)
        w.add_outlay_table(cx.outlay, y)
        w.add_empty_line()

        # Add outlays for thematic.
        for i, th in enumerate(cx.thematics):
            w.add_paragraph(f'3.{i + 1}. ИТОГОВАЯ СМЕТА\n'
                            f'на очередной {y} год и плановый период {y + 1} и {y + 2} годов\n'
                            f'по тематике исследований {th.title}',
                            WD_PARAGRAPH_ALIGNMENT.CENTER, True)
            w.add_outlay_table(th.outlay, y)
            w.add_empty_line()

        w.add_signatures([cx.manager, wsl.shabanov_bm, wsl.smirnnova_oe, wsl.petrischev_av])
        w.save(out + '.docx')

    #-----------------------------------------------------------------------------------------------
    # Order 3188, form, supplement 10.
    #             exec, supplemment 4.
    # Temporary team.
    #-----------------------------------------------------------------------------------------------

    def add_temporary_team_table_with_workload(self, cx, w, year):
        """
        Add temporary team table with workload.

        Parameters
        ----------
        cx : ComplexTheme
            Complex theme.
        w : Worksheet
            Worksheet.
        y : int
            Year.
        """

        # number of people
        n = len(w.lines)
        thn = len(cx.thematics)

        self.add_paragraph('человек/месяц', WD_PARAGRAPH_ALIGNMENT.RIGHT)

        # table and its style
        t = self.doc.add_table(rows=n + 3, cols=7 + thn)
        t.style = 'Table Grid'

        # head
        h = t.rows[0].cells
        h[0].text = '№ п/п'
        h[1].text = 'ФИО'
        h[2].text = 'Должность'
        h[3].text = 'Табельный номер'
        h[4].text = 'Год рождения'
        h[5].text = 'Подразделение'
        h[6].text = f'Планируемые трудозатраты по тематикам исследований '\
                    f'в рамках подтем комплексной темы на {year} год'
        h = t.rows[1].cells
        h[6].text = 'Всего'
        h[7].text = f'Тематика {cx.thematics[0].title}'
        h[8].text = f'Тематика {cx.thematics[1].title}'
        h[9].text = f'Тематика {cx.thematics[2].title}'

        # Merge cells.
        merge_table_cells_in_row(t, 0, 6, 9)
        for j in range(6):
            merge_table_cells_in_column(t, j, 0, 1)

        tot = [0] * (thn + 1)

        # add rest rows
        for i in range(n):
            h = t.rows[i + 2].cells
            wl = w.lines[i]
            e = wl.employee
            p = e.personal
            h[0].text = str(i + 1) + '.'
            h[1].text = p.surname_name_patronymic()
            h[2].text = wl.job_title.name
            h[3].text = e.tabel
            h[4].text = str(p.year)
            h[5].text = wl.job_place.half_full_name
            x = wl.slot * 12
            xj = [x / 3] * thn
            tot[0] = tot[0] + x
            for j in range(thn):
                tot[1 + j] = tot[1 + j] + xj[j]
            h[6].text = f'{utils.norm_digits(x, 2)}'
            for j in range(thn):
                h[7 + j].text = f'{utils.norm_digits(xj[j], 2)}'

        # Total line.
        h = t.rows[n + 2].cells
        h[1].text = 'ИТОГО:'
        for j in range(len(tot)):
            h[6 + j].text = f'{utils.norm_digits(tot[j], 2)}'
        merge_table_cells_in_row(t, n + 2, 1, 5)

        set_table_text_size(t, 8)

    #-----------------------------------------------------------------------------------------------

    def add_temporary_team_table(self, cx, w):
        """
        Add table for temporary team.

        Parameters
        ----------
        cx : ComplexTheme
            Complex theme.
        w : Worksheet
            Worksheet.
        """

        # number of people
        n = len(w.lines)

        # table and its style
        t = self.doc.add_table(rows = n + 1, cols = 7)
        t.style = 'Table Grid'

        # head
        h = t.rows[0].cells
        h[0].text = '№ п/п'
        h[1].text = 'ФИО'
        h[2].text = 'Должность'
        h[3].text = 'Табельный номер'
        h[4].text = 'Год рождения'
        h[5].text = 'Подразделение'
        h[6].text = 'Наименование подтем/тематик исследований'

        # add rest rows
        for i in range(n):
            r = t.rows[i + 1].cells
            wl = w.lines[i]
            e = wl.employee
            p = e.personal
            r[0].text = str(i + 1) + '.'
            r[1].text = p.surname_name_patronymic()
            r[2].text = wl.job_title.name
            r[3].text = e.tabel
            r[4].text = str(p.year)
            r[5].text = wl.job_place.half_full_name
            r[6].text = cx.all_thematics_titles()

    #-----------------------------------------------------------------------------------------------

    @staticmethod
    def generate_form_gos_assignment_3188_10_team(cx, team, y, out):
        """
        Generate form for gos assignment.
        Supplement 10 (temporary team).

        Parameters
        ----------
        cx : ComplexTheme
            Complex theme.
        team : Worksheet
            Worksheet.
        y : int
            Year.
        out : str
            Out file.
        """

        # Create document.
        w = GeneratorWord()

        # Title.
        w.add_paragraph('СОСТАВ\n'
                        'временного трудового коллектива для выполнения научно-исследовательской '
                        f'работы и планируемые трудозатраты на очередной {y} год '
                        f'и плановый период {y + 1} и {y + 2} годов по комплексной теме {cx.title}',
                        WD_PARAGRAPH_ALIGNMENT.CENTER, True)

        # First year.
        w.add_paragraph('1. Состав временного трудового коллектива и планируемые трудозатраты '
                        f'на первый год планового периода - {y} год',
                        WD_PARAGRAPH_ALIGNMENT.CENTER, True)
        w.add_temporary_team_table_with_workload(cx, team, y)
        w.add_empty_line()

        # Second year.
        w.add_paragraph('2. Состав временного трудового коллектива и планируемые трудозатраты '
                        f'на первый год планового периода - {y + 1} год',
                        WD_PARAGRAPH_ALIGNMENT.CENTER, True)
        w.add_temporary_team_table_with_workload(cx, team, y + 1)
        w.add_empty_line()

        # Third year.
        w.add_paragraph('3. Состав временного трудового коллектива и планируемые трудозатраты '
                        f'на первый год планового периода - {y + 2} год',
                        WD_PARAGRAPH_ALIGNMENT.CENTER, True)
        w.add_temporary_team_table_with_workload(cx, team, y + 2)
        w.add_empty_line()

        # Add signatures and close file.
        w.add_signatures([cx.manager, wsl.shabanov_bm, wsl.smirnnova_oe, wsl.petrischev_av])
        w.save(out + '.docx')

    #-----------------------------------------------------------------------------------------------

    @staticmethod
    def generate_exec_gos_assignment_3188_04_team(cx, team, y, out):
        """
        Generate temporary team document.

        Parameters
        ----------
        cx : ComplexTheme
            Complex theme.
        team : Worksheet
            Team working on this theme.
        y : int
            Start year.
        out : str
            Out file name.
        """

        w = GeneratorWord()

        # Inscription.
        w.add_corner_inscription_supplement_to_order(4)
        w.add_empty_line()

        # Title.
        w.add_paragraph('СОСТАВ\nвременного трудового коллектива для выполнения '
                        'научно-исследовательской работы и планируемые трудозатраты '
                        f'на очередной {y} год и плановый период {y + 1} и {y + 2} годов '
                        f'по комплексной теме {cx.title}',
                        WD_PARAGRAPH_ALIGNMENT.CENTER, True)
        w.add_empty_line()

        # Table.
        w.add_temporary_team_table(cx, team)
        w.add_empty_line()

        # Signatures and save.
        w.add_signatures([cx.manager, wsl.shabanov_bm, wsl.smirnnova_oe, wsl.petrischev_av])
        w.save(out + '.docx')

    #-----------------------------------------------------------------------------------------------
    # Order 3188, form, supplement 11.
    #             exec, supplement 5.
    # Generate equipment methods.
    #-----------------------------------------------------------------------------------------------

    def add_inner_equipment_table(self, cx):
        """
        Add inner equipment table.

        Parameters
        ----------
        cx : ComplexTheme
            Complex theme.
        """

        # Add table title.
        self.add_paragraph('1. Перечень находящихся в оперативном управлении '
                           'НИЦ «Курчатовский институт» объектов особо ценного имущества',
                           alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, is_bold=True)

        # Table and its style.
        t = self.doc.add_table(rows=6, cols=4)
        t.style = 'Table Grid'

        # All thematics.
        all_thematics_text = ''
        for th in cx.thematics:
            if all_thematics_text != '':
                all_thematics_text = all_thematics_text + '\n\n'
            all_thematics_text = all_thematics_text + f'Тематика\n{th.title}'

        # Head.
        h = t.rows[0].cells
        h[0].text = '№ п/п'
        h[1].text = 'Наименование объекта особо ценного движимого имущества'
        h[2].text = 'Местоположение (здание, помещение)'
        h[3].text = 'Наименование подтем и тематик исследований'

        # BDW
        h = t.rows[1].cells
        h[0].text = '1.'
        h[1].text = 'Суперкомпьютер МВС-10П ОП, раздел МВС-10П ОП1 BDW (Broadwell),\n'\
                    f'{cx.bdw_need} узлочасов.'
        h[2].text = 'Москва, Ленинский проспект, 32А.'
        h[3].text = all_thematics_text

        # KNL
        h = t.rows[2].cells
        h[0].text = '2.'
        h[1].text = 'Суперкомпьютер МВС-10П ОП, раздел МВС-10П МП2 KNL (Knights Landing),\n'\
                    f'{cx.knl_need} узлочасов.'
        h[2].text = 'Москва, Ленинский проспект, 32А.'
        h[3].text = all_thematics_text

        # CLK
        h = t.rows[3].cells
        h[0].text = '3.'
        h[1].text = 'Суперкомпьютер МВС-10П ОП, раздел МВС-10П ОП2 CLK (Cascade Lake),\n'\
                    f'{cx.clk_need} узлочасов.'
        h[2].text = 'Москва, Ленинский проспект, 32А.'
        h[3].text = all_thematics_text

        # ICL
        h = t.rows[4].cells
        h[0].text = '4.'
        h[1].text = 'Суперкомпьютер МВС-10П ОП, раздел МВС-10П ОП3 ICL (Ice Lake),\n'\
                    f'{cx.icl_need} узлочасов.'
        h[2].text = 'Москва, Ленинский проспект, 32А.'
        h[3].text = all_thematics_text

        # A100
        h = t.rows[5].cells
        h[0].text = '5.'
        h[1].text = 'Вычислительные узлы на базе двух графических карт NVIDIA A100,\n'\
                    f'{cx.a100_need} узлочасов.'
        h[2].text = 'Москва, Ленинский проспект, 32А.'
        h[3].text = all_thematics_text

        set_table_columns_widths(t, [0.5, 2.0, 1.5, 3.5])

    #-----------------------------------------------------------------------------------------------

    def add_outer_equipment_table(self, cx):
        """
        Add outer equipment table.

        Parameters
        ----------
        cx : ComplexTheme
            Complex theme.
        """

        # Add table title.
        self.add_paragraph('2. Перечень объектов особо ценного движимого имущества, '
                           'которое планируется привлечь на условиях аренды',
                           alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, is_bold=True)

        # Table and its style.
        t = self.doc.add_table(rows=2, cols=4)
        t.style = 'Table Grid'

        # Head.
        h = t.rows[0].cells
        h[0].text = '№ п/п'
        h[1].text = 'Наименование объекта особо ценного движимого имущества'
        h[2].text = 'Местоположение (здание, помещение)'
        h[3].text = 'Наименование подтем и тематик исследований'

        set_table_columns_widths(t, [0.5, 2.0, 1.5, 3.5])

    #-----------------------------------------------------------------------------------------------

    @staticmethod
    def generate_form_gos_assignment_3188_11_equipment(cx, out):
        """
        Generate equipment document.

        Parameters
        ----------
        cx : ComplexTheme
            Complex theme.
        out : str
            Out file name.
        """

        w = GeneratorWord()

        # Title.
        w.add_paragraph('ПЕРЕЧЕНЬ\nобъектов особо ценного движимого имущества, '
                        'используемого в процессе выполнения научно-исследовательской работы '
                        '(основных средств и нематериальных активов, амортизируемых в процессе '
                        f'выполнения работы) по комплексной теме {cx.title}',
                        WD_PARAGRAPH_ALIGNMENT.CENTER, True)
        w.add_empty_line()

        # Inner equipment.
        w.add_inner_equipment_table(cx)
        w.add_empty_line()

        # Outer equipment.
        w.add_outer_equipment_table(cx)
        w.add_empty_line()

        # Signatures and save.
        w.add_signatures([cx.manager, wsl.shabanov_bm])
        w.save(out + '.docx')

    #-----------------------------------------------------------------------------------------------

    @staticmethod
    def generate_exec_gos_assignment_3188_05_equipment(cx, out):
        """
        Generate equipment document.

        Parameters
        ----------
        cx : ComplexTheme
            Complex theme.
        out : str
            Out file name.
        """

        w = GeneratorWord()

        # Inscription.
        w.add_corner_inscription_supplement_to_order(5)
        w.add_empty_line()

        # Title.
        w.add_paragraph('ПЕРЕЧЕНЬ\nобъектов особо ценного движимого имущества, '
                        'используемого в процессе выполнения научно-исследовательской работы '
                        '(основных средств и нематериальных активов, амортизируемых в процессе '
                        f'выполнения работы) по комплексной теме {cx.title}',
                        WD_PARAGRAPH_ALIGNMENT.CENTER, True)
        w.add_empty_line()

        # Inner equipment.
        w.add_inner_equipment_table(cx)
        w.add_empty_line()

        # Outer equipment.
        w.add_outer_equipment_table(cx)
        w.add_empty_line()

        # Signatures and save.
        w.add_signatures([cx.manager, wsl.shabanov_bm])
        w.save(out + '.docx')

    #-----------------------------------------------------------------------------------------------
    # Exec gos assignment.
    # Order.
    #-----------------------------------------------------------------------------------------------

    @staticmethod
    def generate_exec_gos_assignment_3188_order(cx, y, out):
        """
        Generate order.

        Parameters
        ----------
        cx : ComplexTheme
            Complex theme.
        y : int
            Year.
        out : str
            Outfile.
        """

        w = GeneratorWord()
        w.add_empty_lines(8)

        w.add_paragraph('О проведении фундаментальной научно-исследовательской работы\n'
                        f'по комплексной теме {cx.title}',
                        WD_PARAGRAPH_ALIGNMENT.CENTER,True, 14)
        w.add_empty_line()
        #
        p = w.add_paragraph('\tВ целях обеспечения выполнения работ по тематическому плану '
                            'научно-исследовательских и опытно-конструкторских работ, '
                            'выполняемых НИЦ «Курчатовский институт» в рамках выполнения '
                            'государственного задания на оказание государственных услуг '
                            f'на {y} год и плановый период {y + 1} и {y + 2} годов '
                            '(далее соответственно – тематический план, государственное задание), '
                            'а также достижения показателей, характеризующих качество работ '
                            'по государственному заданию, ')
        r = p.add_run('п р и к а з ы в а ю:')
        r.bold = True
        #
        w.add_paragraph('\t1. Провести научно-исследовательскую работу по комплексной теме '
                        f'{cx.title} (далее – НИР).')
        #
        w.add_paragraph('\t2. Определить:')
        #
        pers = cx.manager.employee.personal
        perss = pers.surname('ru')
        persn = pers.name_first_letter('ru')
        persp = pers.patronymic_first_letter('ru')
        w.add_paragraph(f'\t2.1. Руководителем работ по комплексной теме '
                        f'{cx.manager.full_job_title_with_job_place_r()} {perss}а {persn}.{persp}.')
        #
        w.add_paragraph('\t2.2 Ответственным структурным подразделением '
                        'НИЦ «Курчатовский институт» за выполнение работ '
                        'отделение суперкомпьютерных систем и параллельных вычислений.')
        #
        w.add_paragraph('\t3. Утвердить:')
        w.add_paragraph('\t3.1. Техническое задание на выполнение НИР согласно '
                        'приложению № 1 к настоящему приказу.')
        w.add_paragraph('\t3.2. Календарный план на выполнение НИР согласно '
                        'приложению № 2 к настоящему приказу.')
        w.add_paragraph('\t3.3. Смету на выполнение НИР согласно '
                        'приложению № 3 к настоящему приказу.')
        w.add_paragraph('\t3.4. Состав временного трудового коллектива для выполнения НИР согласно '
                        'приложению № 4 к настоящему приказу.')
        w.add_paragraph('\t3.5. Перечень объектов особо ценного движимого имущества, используемого '
                        'в процессе выполнения НИР (основных средств и нематериальных активов, '
                        'амортизируемых в процессе выполнения работы) согласно '
                        'приложению № 5 к настоящему приказу.')
        w.add_paragraph(f'\t4. Установить срок выполнения НИР не позднее 31 декабря {y + 2} года. '
                        'Срок выполнения работ по настоящему приказу – '
                        f'не позднее 31 декабря {y} года.')
        w.add_paragraph('\t5. Руководителю работ по комплексной теме, указанному в '
                        'пункте 2.1 настоящего приказа, обеспечить:')
        w.add_paragraph('\tкоординацию выполнения НИР;')
        w.add_paragraph('\tконтроль соблюдения требований и показателей, характеризующих '
                        'качество работы, установленных техническим заданием и '
                        'календарным планом выполнения НИР;')
        w.add_paragraph('\tэффективное и целевое расходование средств, предусмотренных на '
                        'финансовое обеспечение государственного задания '
                        'НИЦ «Курчатовский институт» на выполнение работ в '
                        'пределах доведенных Центру лимитов бюджетных обязательств, '
                        'предусмотренных на выполнение НИР в соответствии со сметой расходов '
                        'на выполнение НИР;')
        w.add_paragraph('\tподготовку единого отчета о НИР, оформленного в соответствии с '
                        'ГОСТ 7.32-2017 и проектов Форм направления сведений для размещения в '
                        'ЕГИСУ НИОКТР;')
        w.add_paragraph('\tполучение положительного заключения профильного экспертного совета '
                        'при Ученом совете НИЦ «Курчатовкий институт» на единый отчет о НИР;')
        w.add_paragraph('\tпредоставление заместителю директора – главному ученому секретарю '
                        f'в срок до 15 января {y + 1} года единого отчета о НИР, '
                        'оформленного в соответствии с приказом Центра от 10 июня 2024 года № 1792 '
                        '(в ред. от 15 августа 2025 года № 3161) «Об организации работы по '
                        'предоставлению сведений о научно-исследовательских, '
                        'опытно-конструкторских и технологических работах, выполняемых '
                        'НИЦ «Курчатовский институт», в единую государственную информационную '
                        'систему учета научно-исследовательских, опытно-конструкторских и '
                        'технологических работ гражданского назначения и в систему «БИТ.Наука» '
                        'на платформе «1СПредприятие» на период опытной эксплуатации '
                        'системы «БИТ.Наука».')
        w.add_paragraph('\t6. Заместителю директора – главному ученому секретарю '
                        f'в срок до 1 февраля {y + 1} г. обеспечить размещение отчета о НИР, '
                        'Форм направления сведений в ЕГИСУ НИОКТР и направление на экспертизу '
                        'в соответствии с Правилами осуществления федеральным государственным '
                        'бюджетным учреждением «Российская академия наук» научного и '
                        'научно-методического руководства научной и научно-технической '
                        'деятельностью научных организаций и образовательных организаций '
                        'высшего образования, а также экспертизы научных и научно-технических '
                        'результатов, полученных этими организациями, утвержденными '
                        'постановлением Правительства Российской Федерации '
                        'от 30 декабря 2018 года № 1781.')
        w.add_paragraph('\t7. Контроль исполнения настоящего приказа оставляю за собой.')
        w.add_empty_line()

        # Signatures and save.
        w.add_signatures([wsl.dyakova_ya], False)
        w.save(out + '.docx')

#===================================================================================================

if __name__ == '__main__':
    pass

#===================================================================================================
